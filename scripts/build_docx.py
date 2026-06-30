#!/usr/bin/env python3
"""
build_docx.py — render docs/manuscript_draft.md into a formatted, figure-embedded
Word document docs/manuscript_draft.docx (US Letter), faithfully (content
unchanged; only formatting + embedded images).

Five figures embedded from figures/ at their captions:
  Fig 1 -> cross_section_comparison.png
  Fig 2 -> sensitivity_2d.png
  Fig 3 -> montecarlo_hist.png + montecarlo_projection.png
  Fig 4 -> pulsed_timeseries.png + pulsed_Gmap.png
  Fig 5 -> radial_profiles.png + radial_Pnet_maps.png
"""
import os
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.normpath(os.path.join(_HERE, '..'))
MD = os.path.join(ROOT, 'docs', 'manuscript_draft.md')
OUT = os.path.join(ROOT, 'docs', 'manuscript_draft.docx')
FIGDIR = os.path.join(ROOT, 'figures')

FIG_MAP = {
    '1': ['cross_section_comparison.png'],
    '2': ['sensitivity_2d.png'],
    '3': ['montecarlo_hist.png', 'montecarlo_projection.png'],
    '4': ['pulsed_timeseries.png', 'pulsed_Gmap.png'],
    '5': ['radial_profiles.png', 'radial_Pnet_maps.png'],
}

BODY_FONT = 'Times New Roman'
CODE_FONT = 'Consolas'
IMG_W = Inches(6.4)

# subscript: base char (latin/greek/⟩/∂) + _ + {...} or word
SUB_RE = re.compile(r'([A-Za-zͰ-ω⟩∂])_(\{[^}]*\}|[A-Za-z0-9Ͱ-ω]+)')
# superscript: ^ then {...} / (...) / token (incl greek, minus sign) / ∞
SUP_RE = re.compile(r'\^(\{[^}]*\}|\([^)]*\)|[A-Za-z0-9.−Ͱ-ω\-+/]+|∞)')
# inline emphasis / code
SPAN_RE = re.compile(r'(`[^`]*`|\*\*[^*]+\*\*|\*[^*]+\*)')


def _scriptify(s):
    """Yield (text, script) where script in {None,'sub','sup'}."""
    out, buf, i, n = [], '', 0, len(s)
    while i < n:
        msub = SUB_RE.match(s, i)
        msup = SUP_RE.match(s, i)
        if msub:
            if buf:
                out.append((buf, None)); buf = ''
            out.append((msub.group(1), None))
            out.append((msub.group(2).strip('{}'), 'sub'))
            i = msub.end()
        elif msup:
            if buf:
                out.append((buf, None)); buf = ''
            g = msup.group(1)
            if g and g[0] in '{(':
                g = g.strip('{}()')
            out.append((g, 'sup'))
            i = msup.end()
        else:
            buf += s[i]; i += 1
    if buf:
        out.append((buf, None))
    return out


def _emit(p, text, bold=False, italic=False, code=False, size=None, color=None):
    if code:
        r = p.add_run(text)
        r.font.name = CODE_FONT
        r.font.size = Pt((size or 11) - 1)
        if color:
            r.font.color.rgb = color
        return
    for seg, script in _scriptify(text):
        if seg == '':
            continue
        r = p.add_run(seg)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        if script == 'sub':
            r.font.subscript = True
        elif script == 'sup':
            r.font.superscript = True


def render_inline(p, text, size=None, color=None):
    for piece in SPAN_RE.split(text):
        if not piece:
            continue
        if piece.startswith('`') and piece.endswith('`'):
            _emit(p, piece[1:-1], code=True, size=size, color=color)
        elif piece.startswith('**') and piece.endswith('**'):
            _emit(p, piece[2:-2], bold=True, size=size, color=color)
        elif piece.startswith('*') and piece.endswith('*') and len(piece) > 2:
            _emit(p, piece[1:-1], italic=True, size=size, color=color)
        else:
            _emit(p, piece, size=size, color=color)


def shade(cell_or_para, hexfill):
    el = cell_or_para
    pr = el.get_or_add_tcPr() if hasattr(el, 'get_or_add_tcPr') else el.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexfill)
    pr.append(sh)


def add_heading(doc, text, size, before=12, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return p


def add_table(doc, rows):
    header = [c.strip() for c in rows[0].strip('|').split('|')]
    sep = [c.strip() for c in rows[1].strip('|').split('|')]
    aligns = []
    for s in sep:
        aligns.append(WD_ALIGN_PARAGRAPH.RIGHT if s.endswith(':') and not s.startswith(':')
                      else (WD_ALIGN_PARAGRAPH.CENTER if s.startswith(':') and s.endswith(':')
                            else WD_ALIGN_PARAGRAPH.LEFT))
    data = [[c.strip() for c in r.strip('|').split('|')] for r in rows[2:]]
    ncol = len(header)
    t = doc.add_table(rows=1 + len(data), cols=ncol)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        c.paragraphs[0].clear()
        render_inline(c.paragraphs[0], h, size=10)
        for run in c.paragraphs[0].runs:
            run.bold = True
        c.paragraphs[0].alignment = aligns[j] if j < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
        shade(c._tc, 'D9E5F1')
    for i, drow in enumerate(data):
        for j in range(ncol):
            c = t.rows[i + 1].cells[j]
            c.paragraphs[0].clear()
            txt = drow[j] if j < len(drow) else ''
            render_inline(c.paragraphs[0], txt, size=10)
            c.paragraphs[0].alignment = aligns[j] if j < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_figure_images(doc, num):
    for fn in FIG_MAP.get(num, []):
        path = os.path.join(FIGDIR, fn)
        if os.path.exists(path):
            doc.add_picture(path, width=IMG_W)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].paragraph_format.space_before = Pt(6)


def main():
    lines = open(MD, encoding='utf-8').read().split('\n')
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)
    st = doc.styles['Normal']
    st.font.name = BODY_FONT
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15

    frontmatter = True
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].rstrip()
        s = line.strip()

        if s.startswith('## '):
            frontmatter = False

        if frontmatter:
            if s.startswith('# '):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(s[2:]); r.bold = True; r.font.size = Pt(15)
            elif s == '---' or s == '':
                pass
            elif s.startswith('**') and s.endswith('**'):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _emit(p, s[2:-2], bold=True, size=13)
            elif s.startswith('*') and s.endswith('*'):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _emit(p, s[1:-1], italic=True, size=9.5)
            else:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                render_inline(p, s, size=10)
            i += 1
            continue

        if s == '' or s == '---':
            i += 1
            continue
        if s.startswith('## '):
            add_heading(doc, s[3:], 14, before=14, after=6)
            i += 1
            continue
        if s.startswith('### '):
            add_heading(doc, s[4:], 12, before=10, after=4)
            i += 1
            continue
        if s.startswith('|'):
            block = []
            while i < n and lines[i].strip().startswith('|'):
                block.append(lines[i].strip()); i += 1
            if len(block) >= 2:
                add_table(doc, block)
            continue
        if s.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            shade(p._p, 'F2F2F2')
            render_inline(p, s[2:], size=10)
            i += 1
            continue
        mfig = re.match(r'\*\*Figure (\d+)\.\*\*', s)
        if mfig:
            add_figure_images(doc, mfig.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            render_inline(p, s, size=10)
            i += 1
            continue
        if s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            render_inline(p, s[2:])
            i += 1
            continue
        mref = re.match(r'^(\d+)\.\s+(.*)', s)
        if mref:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            p.paragraph_format.space_after = Pt(2)
            _emit(p, mref.group(1) + '. ', size=10)
            render_inline(p, mref.group(2), size=10)
            i += 1
            continue
        # default paragraph
        p = doc.add_paragraph()
        render_inline(p, s)
        i += 1

    doc.save(OUT)
    print('[OK] wrote', OUT)


if __name__ == '__main__':
    main()
